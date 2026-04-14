import csv
import io
import datetime
from typing import List
import pandas as pd

from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.barcharts import VerticalBarChart

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -- REPORTES DE VENTAS --

def generar_pdf_ventas(ventas: List[dict], usuario: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    elements = []
    styles = getSampleStyleSheet()
    
    # Estilos Personalizados
    title_style = ParagraphStyle(
        'TitleStyle', parent=styles['Heading1'],
        fontName='Helvetica-Bold', fontSize=18, textColor=colors.HexColor('#1E3A8A'),
        alignment=1, spaceAfter=10
    )
    subtitle_style = ParagraphStyle(
        'SubTitle', parent=styles['Normal'],
        fontName='Helvetica', fontSize=10, textColor=colors.gray, alignment=1
    )
    
    # Encabezado
    fecha_hoy = datetime.datetime.now().strftime("%d/%m/%Y")
    rango_fechas = f"{ventas[-1]['fecha']} al {ventas[0]['fecha']}" if ventas else "Sin datos"
    
    elements.append(Paragraph("MACUIN - Reporte Ejecutivo de Ventas", title_style))
    elements.append(Paragraph(f"<b>Sucursal:</b> Matriz Central  |  <b>Generado por:</b> {usuario}  |  <b>Fecha:</b> {fecha_hoy}", subtitle_style))
    elements.append(Paragraph(f"<b>Periodo:</b> {rango_fechas}", subtitle_style))
    elements.append(Spacer(1, 0.2 * inch))
    
    # Datos de Ventas
    table_data = [["Fecha", "ID Venta", "Cliente", "Categoría", "Subtotal", "IVA", "Total"]]
    tot_ingresos = 0
    ventas_diarias = {}
    
    for v in ventas:
        table_data.append([
            v['fecha'], str(v['id_venta']), v['cliente'][:20], v['categoria'][:15],
            f"${v['subtotal']:,.2f}", f"${v['iva']:,.2f}", f"${v['total']:,.2f}"
        ])
        tot_ingresos += v['total']
        ventas_diarias[v['fecha']] = ventas_diarias.get(v['fecha'], 0) + v['total']
        
    table_data.append(["", "", "", "TOTALES:", "", "", f"${tot_ingresos:,.2f}"])
    
    # Estilo de Tabla
    table = Table(table_data, colWidths=[65, 55, 120, 100, 70, 70, 70])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('BACKGROUND', (0, 1), (-1, -2), colors.white),
        ('GRID', (0, 0), (-1, -2), 0.5, colors.lightgrey),
        ('FONTNAME', (-3, -1), (-1, -1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (-1, -1), (-1, -1), colors.HexColor('#10B981')), # Green total
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 0.4 * inch))
    
    # Gráfico de Barras Estilizado
    if ventas_diarias:
        d = Drawing(400, 200)
        chart = VerticalBarChart()
        chart.x = 50
        chart.y = 50
        chart.height = 125
        chart.width = 300
        
        datos = list(ventas_diarias.items())
        datos.sort(key=lambda x: x[0])
        fechas = [f.split('-')[-1] for f, _ in datos][-10:] # Last 10 days
        valores = [v for _, v in datos][-10:]
        
        chart.data = [valores]
        chart.categoryAxis.categoryNames = fechas
        chart.bars[0].fillColor = colors.HexColor('#3B82F6')
        chart.valueAxis.valueMin = 0
        
        d.add(chart)
        elements.append(Paragraph("<b>Tendencia de Ventas (Últimos días)</b>", styles['Normal']))
        elements.append(d)

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generar_xlsx_ventas(ventas: List[dict], usuario: str) -> io.BytesIO:
    buffer = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "Ventas por Periodo"
    
    # Estilos Header
    header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    headers = ["Fecha", "ID Venta", "Cliente", "Categoría", "Subtotal", "IVA", "Total"]
    ws.append(headers)
    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        
    for i, v in enumerate(ventas, start=2):
        ws.append([
            v['fecha'], v['id_venta'], v['cliente'], v['categoria'],
            v['subtotal'], v['iva'], v['total']
        ])
        # Formato de Moneda
        ws.cell(row=i, column=5).number_format = '"$"#,##0.00'
        ws.cell(row=i, column=6).number_format = '"$"#,##0.00'
        ws.cell(row=i, column=7).number_format = '"$"#,##0.00'

    # Fila de Totales Dinámica
    ult_fila = len(ventas) + 1
    ws.cell(row=ult_fila+1, column=4, value="TOTALES:").font = Font(bold=True)
    ws.cell(row=ult_fila+1, column=5, value=f"=SUM(E2:E{ult_fila})").number_format = '"$"#,##0.00'
    ws.cell(row=ult_fila+1, column=6, value=f"=SUM(F2:F{ult_fila})").number_format = '"$"#,##0.00'
    ws.cell(row=ult_fila+1, column=7, value=f"=SUM(G2:G{ult_fila})").number_format = '"$"#,##0.00'
    ws.cell(row=ult_fila+1, column=7).font = Font(bold=True, color="10B981")

    # Autofilter
    ws.auto_filter.ref = f"A1:G{ult_fila}"
    
    # Ancho de columnas
    for col in range(1, 8):
        ws.column_dimensions[get_column_letter(col)].width = 15

    wb.save(buffer)
    buffer.seek(0)
    return buffer

def generar_docx_ventas(ventas: List[dict], usuario: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = Document()
    
    # Título corporativo
    title = doc.add_heading('Informe Ejecutivo de Rentabilidad y Ventas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f"Generado por: {usuario}\n").bold = True
    p.add_run(f"Sucursal: MACUIN Matriz Central\n")
    p.add_run(f"Fecha de Creación: {datetime.datetime.now().strftime('%d/%m/%Y')}\n")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph("Resumen Financiero", style='Heading 2')
    tot_ingresos = sum(v['total'] for v in ventas)
    doc.add_paragraph(f"Durante el periodo evaluado, la sucursal generó ingresos totales equivalentes a ${tot_ingresos:,.2f} MXN.")
    
    doc.add_paragraph("Desglose de Ventas", style='Heading 2')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fecha'
    hdr_cells[1].text = 'ID Venta'
    hdr_cells[2].text = 'Cliente'
    hdr_cells[3].text = 'Total'
    
    for v in ventas:
        row_cells = table.add_row().cells
        row_cells[0].text = v['fecha']
        row_cells[1].text = str(v['id_venta'])
        row_cells[2].text = v['cliente'][:20]
        row_cells[3].text = f"${v['total']:,.2f}"
        
    doc.add_paragraph("\nConclusiones y Observaciones:", style='Heading 2')
    doc.add_paragraph("[Escriba aquí el análisis de ventas, desviaciones de rentabilidad o notas logísticas...]")
    
    doc.add_paragraph("\n\n\n\n____________________________________", style='Normal')
    doc.add_paragraph("Firma de Autorización - Gerencia", style='Normal')
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -- REPORTES DE INVENTARIO --

def generar_pdf_inventario(inventario: List[dict]) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontName='Helvetica-Bold', textColor=colors.HexColor('#F59E0B'), alignment=1)
    elements.append(Paragraph("Reporte de Control de Stock", title_style))
    elements.append(Spacer(1, 0.2 * inch))
    
    table_data = [["SKU", "Familia", "Descripción", "Marca", "Stock Actual", "Reorden", "Valorizado"]]
    
    for item in inventario:
        row = [
            item["codigo_parte"], item["familia"], item["descripcion"][:25], item["marca"],
            str(item["stock_actual"]), str(item["punto_reorden"]), f"${item['valorizado']:,.2f}"
        ]
        table_data.append(row)
        
    table = Table(table_data, colWidths=[60, 80, 160, 80, 70, 70, 80])
    
    # Resaltar en ROJO si stock < reorden
    estilos = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F59E0B')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
    ]
    
    for i, item in enumerate(inventario, start=1):
        if item["stock_actual"] < item["punto_reorden"]:
            estilos.append(('TEXTCOLOR', (4, i), (4, i), colors.red))
            estilos.append(('FONTNAME', (4, i), (4, i), 'Helvetica-Bold'))
            
    table.setStyle(TableStyle(estilos))
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generar_xlsx_inventario(inventario: List[dict]) -> io.BytesIO:
    buffer = io.BytesIO()
    df = pd.DataFrame(inventario)
    
    # Renombrar columnas para estetica
    df = df.rename(columns={
        "codigo_parte": "SKU", "marca": "Marca", "descripcion": "Descripción", 
        "stock_actual": "Stock Actual", "punto_reorden": "Punto Reorden", 
        "valorizado": "Valorizado ($)", "familia": "Familia"
    })
    # Reordenar
    df = df[["SKU", "Familia", "Descripción", "Marca", "Stock Actual", "Punto Reorden", "Valorizado ($)"]]
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Control de Stock"
    
    from openpyxl.utils.dataframe import dataframe_to_rows
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import PatternFill
    
    for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
        ws.append(row)
        
    # Colorear Header
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="F59E0B")
        
    # Formato numérico
    for row in range(2, len(df)+2):
        ws.cell(row=row, column=7).number_format = '"$"#,##0.00'
        
    # Condicional para quiebre de stock (Si Stock Actual < Reorden)
    # OpenPyXL es complejo comparando dos celdas, colorearé manualmente
    red_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
    red_font = Font(color='FF0000', bold=True)
    
    for i in range(2, len(df)+2):
        stock_actual = ws.cell(row=i, column=5).value
        reorden = ws.cell(row=i, column=6).value
        if stock_actual is not None and reorden is not None and stock_actual < reorden:
            ws.cell(row=i, column=5).fill = red_fill
            ws.cell(row=i, column=5).font = red_font
            
    ws.auto_filter.ref = ws.dimensions
    for col in range(1, 8):
        ws.column_dimensions[get_column_letter(col)].width = 15
        
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def generar_docx_inventario(inventario: List[dict]) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = Document()
    title = doc.add_heading('Informe de Control de Stock y Quiebres', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    criticos = [i for i in inventario if i["stock_actual"] < i["punto_reorden"]]
    doc.add_paragraph(f"Resumen Ejecutivo:", style='Heading 2')
    doc.add_paragraph(f"Se identificaron {len(criticos)} piezas críticas por debajo del punto de reorden.")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Medium Shading 1 Accent 3'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SKU'
    hdr_cells[1].text = 'Descripción'
    hdr_cells[2].text = 'Stock / Reorden'
    hdr_cells[3].text = 'Familia'
    
    for item in inventario:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["codigo_parte"])
        row_cells[1].text = item["descripcion"][:20]
        st = f"{item['stock_actual']} / {item['punto_reorden']}"
        row_cells[2].text = st
        if item["stock_actual"] < item["punto_reorden"]:
            row_cells[2].paragraphs[0].runs[0].font.bold = True
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
        row_cells[3].text = item["familia"]
        
    doc.add_paragraph("\nComentarios de Almacenaje:", style='Heading 2')
    doc.add_paragraph("________________________________________________________\n" * 3)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -- REPORTES DE PEDIDOS Y LOGÍSTICA --

def generar_pdf_pedidos(pedidos: List[dict]) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontName='Helvetica-Bold', textColor=colors.HexColor('#3B82F6'), alignment=1)
    elements.append(Paragraph("Logística: Gestión de Pedidos y Distribución", title_style))
    elements.append(Spacer(1, 0.2 * inch))
    
    for p in pedidos:
        # Ficha de Pedido
        p_text = f"<b>📌 Pedido #{p['id']}</b> | Fecha: {p['fecha']} | Cliente: {p['cliente']} | Estado: <font color='{'red' if p['estado']=='CANCELADO' else 'green'}'>{p['estado']}</font>"
        elements.append(Paragraph(p_text, styles['Normal']))
        
        # Detalle de piezas
        elements.append(Paragraph(f"<i>Piezas:</i> {p['piezas']}", styles['Italic']))
        elements.append(Spacer(1, 0.1 * inch))
        
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generar_xlsx_pedidos(pedidos: List[dict]) -> io.BytesIO:
    buffer = io.BytesIO()
    df = pd.DataFrame(pedidos)
    df = df.rename(columns={"id": "ID Pedido", "cliente": "Cliente", "fecha": "Fecha", "estado": "Estado de Surtido", "piezas": "Detalle de Piezas"})
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Pedidos y Logística"
    
    from openpyxl.utils.dataframe import dataframe_to_rows
    from openpyxl.styles import PatternFill
    for row in dataframe_to_rows(df, index=False, header=True):
        ws.append(row)
        
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="3B82F6")
        
    ws.auto_filter.ref = ws.dimensions
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['E'].width = 50
    
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def generar_docx_pedidos(pedidos: List[dict]) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = Document()
    title = doc.add_heading('Expediente Logístico y de Distribución', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for p in pedidos:
        doc.add_heading(f"Pedido #{p['id']} - {p['cliente']}", level=2)
        doc.add_paragraph(f"Fecha de Solicitud: {p['fecha']}")
        doc.add_paragraph(f"Estatus de Entrega: {p['estado']}")
        doc.add_paragraph(f"Detalle de Empaque:")
        doc.add_paragraph(p['piezas'], style='List Bullet')
        doc.add_paragraph("")

    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -- CSV genérico --

def _csv_bytes(rows: List[dict]) -> io.BytesIO:
    if not rows:
        return io.BytesIO('Sin datos\n'.encode('utf-8-sig'))
    out = io.StringIO()
    writer = csv.DictWriter(out, fieldnames=rows[0].keys())
    writer.writeheader()
    writer.writerows(rows)
    return io.BytesIO(out.getvalue().encode('utf-8-sig'))


def generar_csv_ventas(ventas: List[dict]) -> io.BytesIO:
    return _csv_bytes(ventas)


def generar_csv_pedidos(pedidos: List[dict]) -> io.BytesIO:
    return _csv_bytes(pedidos)


def generar_csv_inventario(inventario: List[dict]) -> io.BytesIO:
    return _csv_bytes(inventario)


def generar_csv_clientes(clientes: List[dict]) -> io.BytesIO:
    return _csv_bytes(clientes)


# -- REPORTE DE CLIENTES --

def generar_pdf_clientes(clientes: List[dict], usuario: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'Title', parent=styles['Heading1'],
        fontName='Helvetica-Bold', fontSize=16,
        textColor=colors.HexColor('#1E3A8A'), alignment=1,
    )
    sub_style = ParagraphStyle(
        'Sub', parent=styles['Normal'],
        fontName='Helvetica', fontSize=9,
        textColor=colors.gray, alignment=1,
    )

    fecha_hoy = datetime.datetime.now().strftime('%d/%m/%Y')
    elements.append(Paragraph('MACUIN - Reporte de Clientes', title_style))
    elements.append(Paragraph(
        f'Generado por: {usuario}  |  Fecha: {fecha_hoy}', sub_style,
    ))
    elements.append(Spacer(1, 0.3 * inch))

    table_data = [['#', 'Cliente', 'Total Pedidos', 'Total Gastado']]
    for i, c in enumerate(clientes, 1):
        table_data.append([
            str(i), c['nombre'], str(c['total_pedidos']),
            f"${c['total_gastado']:,.2f}",
        ])

    table = Table(table_data, colWidths=[30, 200, 100, 120])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
    ]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generar_xlsx_clientes(clientes: List[dict], usuario: str) -> io.BytesIO:
    buffer = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = 'Clientes'

    header_fill = PatternFill(start_color='1E3A8A', end_color='1E3A8A', fill_type='solid')
    header_font = Font(color='FFFFFF', bold=True)

    headers = ['#', 'Cliente', 'Total Pedidos', 'Total Gastado']
    ws.append(headers)
    for col in range(1, 5):
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')

    for i, c in enumerate(clientes, 1):
        ws.append([i, c['nombre'], c['total_pedidos'], c['total_gastado']])
        ws.cell(row=i + 1, column=4).number_format = '"$"#,##0.00'

    for col_letter, width in [('A', 6), ('B', 30), ('C', 16), ('D', 18)]:
        ws.column_dimensions[col_letter].width = width

    wb.save(buffer)
    buffer.seek(0)
    return buffer


def generar_docx_clientes(clientes: List[dict], usuario: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = Document()
    title = doc.add_heading('Reporte de Clientes - MACUIN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f'Generado por: {usuario}\n').bold = True
    p.add_run(f'Fecha: {datetime.datetime.now().strftime("%d/%m/%Y")}\n')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph('Ranking por Volumen de Compra', style='Heading 2')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '#'
    hdr[1].text = 'Cliente'
    hdr[2].text = 'Pedidos'
    hdr[3].text = 'Total Gastado'

    for i, c in enumerate(clientes, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = c['nombre']
        row[2].text = str(c['total_pedidos'])
        row[3].text = f"${c['total_gastado']:,.2f}"

    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── Factura por Pedido ─────────────────────────────────────────────────────────

def generar_pdf_factura(pedido: dict) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch,
    )
    elements = []
    styles = getSampleStyleSheet()

    AZUL  = colors.HexColor('#1E3A8A')
    ROJO  = colors.HexColor('#DC2626')
    GRIS  = colors.HexColor('#F1F5F9')
    NEGRO = colors.HexColor('#1E293B')

    title_style = ParagraphStyle(
        'FacTitle', fontName='Helvetica-Bold', fontSize=22,
        textColor=AZUL, spaceAfter=2,
    )
    sub_style = ParagraphStyle(
        'FacSub', fontName='Helvetica', fontSize=9,
        textColor=colors.HexColor('#64748B'), spaceAfter=2,
    )
    label_style = ParagraphStyle(
        'FacLabel', fontName='Helvetica-Bold', fontSize=8,
        textColor=colors.HexColor('#64748B'), spaceAfter=1,
    )
    value_style = ParagraphStyle(
        'FacValue', fontName='Helvetica', fontSize=9,
        textColor=NEGRO, spaceAfter=2,
    )
    footer_style = ParagraphStyle(
        'FacFooter', fontName='Helvetica', fontSize=7,
        textColor=colors.gray, alignment=1,
    )

    pedido_id   = pedido.get('id', 0)
    fecha_raw   = pedido.get('fecha_pedido', '')
    try:
        fecha_dt = datetime.datetime.fromisoformat(str(fecha_raw)[:19])
        fecha_str = fecha_dt.strftime('%d/%m/%Y')
    except Exception:
        fecha_str = str(fecha_raw)[:10]

    cliente      = pedido.get('usuario', {}).get('nombre', 'Cliente')
    direccion    = pedido.get('direccion_entrega') or 'No especificada'
    estado       = pedido.get('estado', 'RECIBIDO')
    subtotal    = float(pedido.get('subtotal', 0))
    envio       = 0.0 if subtotal >= 1000 else round(subtotal * 0.15, 2)
    iva         = round((subtotal + envio) * 0.16, 2)
    total_final = round((subtotal + envio) * 1.16, 2)
    folio        = f"FAC-{pedido_id:05d}"
    fecha_hoy    = datetime.datetime.now().strftime('%d/%m/%Y')

    # ── Cabecera ────────────────────────────────────────────────────────────────
    header_data = [[
        Paragraph("<b>MACUIN</b>", title_style),
        Paragraph(f"<b>FACTURA</b>", ParagraphStyle(
            'FacTitleR', fontName='Helvetica-Bold', fontSize=22,
            textColor=ROJO, alignment=2,
        )),
    ]]
    header_sub = [[
        Paragraph("Refacciones Automotrices<br/>contacto@macuin.com", sub_style),
        Paragraph(
            f"<b>Folio:</b> {folio}<br/>"
            f"<b>Fecha de emisión:</b> {fecha_hoy}<br/>"
            f"<b>Fecha de pedido:</b> {fecha_str}",
            ParagraphStyle('FacSubR', fontName='Helvetica', fontSize=9,
                           textColor=colors.HexColor('#64748B'), alignment=2),
        ),
    ]]

    for row in [header_data, header_sub]:
        t = Table(row, colWidths=[3.5*inch, 3.5*inch])
        t.setStyle(TableStyle([
            ('VALIGN',     (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]))
        elements.append(t)

    elements.append(Spacer(1, 0.15*inch))

    # Línea separadora
    sep = Table([['']], colWidths=[7*inch])
    sep.setStyle(TableStyle([
        ('LINEABOVE', (0,0), (-1,-1), 2, AZUL),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
    ]))
    elements.append(sep)
    elements.append(Spacer(1, 0.15*inch))

    # ── Datos del cliente ───────────────────────────────────────────────────────
    cliente_data = [[
        Paragraph("FACTURAR A", ParagraphStyle(
            'SecTitle', fontName='Helvetica-Bold', fontSize=8,
            textColor=colors.white, spaceAfter=0,
        )),
        Paragraph("ESTADO DEL PEDIDO", ParagraphStyle(
            'SecTitle2', fontName='Helvetica-Bold', fontSize=8,
            textColor=colors.white, spaceAfter=0,
        )),
    ]]
    t = Table(cliente_data, colWidths=[3.5*inch, 3.5*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), AZUL),
        ('TEXTCOLOR',     (0,0), (-1,-1), colors.white),
        ('TOPPADDING',    (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING',   (0,0), (-1,-1), 8),
    ]))
    elements.append(t)

    estado_color_map = {
        'RECIBIDO': '#F59E0B', 'SURTIDO': '#F97316',
        'ENVIADO':  '#16A34A', 'CANCELADO': '#6B7280',
    }
    estado_color = colors.HexColor(estado_color_map.get(estado, '#6B7280'))

    info_data = [[
        Paragraph(f"{cliente}<br/><font size='8' color='#64748B'>Dirección: {direccion}</font>", value_style),
        Paragraph(
            f"<font color='{estado_color_map.get(estado, '#6B7280')}'><b>{estado}</b></font>",
            ParagraphStyle('EstadoStyle', fontName='Helvetica-Bold', fontSize=11,
                           textColor=estado_color, alignment=0, leftIndent=8),
        ),
    ]]
    t = Table(info_data, colWidths=[3.5*inch, 3.5*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), GRIS),
        ('TOPPADDING',    (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('LEFTPADDING',   (0,0), (-1,-1), 8),
        ('BOX',           (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.2*inch))

    # ── Tabla de partidas ───────────────────────────────────────────────────────
    col_widths = [0.4*inch, 2.5*inch, 1.1*inch, 0.7*inch, 1.2*inch, 1.1*inch]

    items_data = [['#', 'Descripción', 'Marca', 'Cant.', 'P. Unitario', 'Importe']]
    detalles = pedido.get('detalles', [])
    for i, d in enumerate(detalles, start=1):
        auto    = d.get('autoparte', {}) or {}
        cant    = d.get('cantidad', 0)
        precio  = float(d.get('precio_unitario', 0))
        importe = cant * precio
        items_data.append([
            str(i),
            auto.get('nombre', '—'),
            auto.get('marca', '—') or '—',
            str(cant),
            f"${precio:,.2f}",
            f"${importe:,.2f}",
        ])

    t = Table(items_data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,0), AZUL),
        ('TEXTCOLOR',     (0,0), (-1,0), colors.white),
        ('FONTNAME',      (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',      (0,0), (-1,0), 8),
        ('TOPPADDING',    (0,0), (-1,0), 7),
        ('BOTTOMPADDING', (0,0), (-1,0), 7),
        ('ALIGN',         (0,0), (-1,0), 'CENTER'),
        ('FONTNAME',      (0,1), (-1,-1), 'Helvetica'),
        ('FONTSIZE',      (0,1), (-1,-1), 8),
        ('TOPPADDING',    (0,1), (-1,-1), 6),
        ('BOTTOMPADDING', (0,1), (-1,-1), 6),
        ('ALIGN',         (0,1), (0,-1), 'CENTER'),
        ('ALIGN',         (3,1), (-1,-1), 'RIGHT'),
        ('ROWBACKGROUNDS',(0,1), (-1,-1), [colors.white, GRIS]),
        ('GRID',          (0,0), (-1,-1), 0.3, colors.HexColor('#E2E8F0')),
        ('LINEBELOW',     (0,0), (-1,0), 1, AZUL),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.15*inch))

    # ── Totales ──────────────────────────────────────────────────────────────────
    envio_label = 'Envío: Gratis' if envio == 0 else f'Envío (15%):'
    envio_valor = '—' if envio == 0 else f"${envio:,.2f}"
    totales_data = [
        ['', 'Subtotal:', f"${subtotal:,.2f}"],
        ['', envio_label, envio_valor],
        ['', 'IVA (16%):', f"${iva:,.2f}"],
        ['', 'TOTAL:', f"${total_final:,.2f}"],
    ]
    t = Table(totales_data, colWidths=[4.3*inch, 1.5*inch, 1.2*inch])
    t.setStyle(TableStyle([
        ('ALIGN',         (1,0),  (-1,-1), 'RIGHT'),
        ('FONTNAME',      (1,0),  (-1,-2), 'Helvetica'),
        ('FONTNAME',      (1,-1), (-1,-1), 'Helvetica-Bold'),
        ('FONTSIZE',      (0,0),  (-1,-1), 9),
        ('FONTSIZE',      (1,-1), (-1,-1), 11),
        ('TEXTCOLOR',     (1,-1), (-1,-1), AZUL),
        ('TOPPADDING',    (0,0),  (-1,-1), 4),
        ('BOTTOMPADDING', (0,0),  (-1,-1), 4),
        ('LINEABOVE',     (1,-1), (-1,-1), 1, AZUL),
        ('BACKGROUND',    (1,-1), (-1,-1), colors.HexColor('#EFF6FF')),
        ('TEXTCOLOR',     (1,1),  (-1,1),  colors.HexColor('#16A34A') if envio == 0 else colors.black),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.3*inch))

    # ── Pie ───────────────────────────────────────────────────────────────────────
    elements.append(sep)
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph(
        "Este documento es una representación impresa de factura. "
        "MACUIN — Refacciones Automotrices. Todos los precios en MXN.",
        footer_style,
    ))

    doc.build(elements)
    buffer.seek(0)
    return buffer
